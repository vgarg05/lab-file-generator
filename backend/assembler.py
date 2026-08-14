"""
assembler.py — Build the final DOCX document using python-docx.

Layout (per spec):
  - Centered heading:  "Experiment / <N>"   — TNR 16pt Bold
  - Section headings:  Aim / Theory / Code / Output  — TNR 16pt Bold, left-aligned
  - Body text:         TNR 14pt, left-aligned
  - Code block:        TNR 14pt, left-aligned, preserves indentation
  - Output section:    terminal PNG image
"""

import io

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Font helper ────────────────────────────────────────────────────────────────

def _apply_tnr(run, size_pt: float, bold: bool = False):
    """Apply Times New Roman to a run, covering ASCII + complex scripts."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold

    # Ensure the font is propagated to all character sets in the XML
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Times New Roman")


# ── Document element builders ──────────────────────────────────────────────────

def _add_section_heading(doc: Document, text: str):
    """Add a left-aligned section heading: TNR 16pt Bold."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text + ":")
    _apply_tnr(run, 16, bold=True)


def _add_body_paragraph(doc: Document, text: str):
    """Add a body text paragraph: TNR 14pt, left-aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text.strip())
    _apply_tnr(run, 14, bold=False)


def _add_inline_section(doc: Document, heading: str, body: str):
    """Add inline heading (16pt Bold) followed by body (14pt Regular) on the same line."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)

    # Heading (16pt Bold) with colon
    h_run = para.add_run(heading + ":")
    _apply_tnr(h_run, 16, bold=True)

    # Body (14pt Regular) — prefixed with a single space
    b_run = para.add_run(" " + body.strip())
    _apply_tnr(b_run, 14, bold=False)


def _add_code_block(doc: Document, code: str):
    """
    Add source code preserving line breaks.
    Uses a single paragraph with explicit line-break runs — TNR 14pt.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)

    lines = code.splitlines()
    for i, line in enumerate(lines):
        run = para.add_run(line)
        _apply_tnr(run, 14, bold=False)
        if i < len(lines) - 1:
            run.add_break()  # Word line-break, stays in one paragraph


# ── Main assembler ─────────────────────────────────────────────────────────────

def assemble_document(
    experiment_number: int,
    aim: str,
    theory: str,
    code: str,
    terminal_png: bytes,
    software: str,
    plot_png: bytes | None = None,
) -> bytes:
    """
    Assemble and return the lab experiment as a DOCX bytes object.

    Args:
        experiment_number: Integer shown in the title ("Experiment / N").
        aim:               Experiment aim string.
        theory:            AI-generated theory text (may contain paragraph breaks).
        code:              Source code string.
        terminal_png:      PNG bytes of the rendered terminal output.
        software:          List of software used.

    Returns:
        DOCX file as raw bytes (in-memory, no temp files written).
    """
    doc = Document()

    # ── Default style ────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    # ── Page: A4, 1-inch margins ─────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Title: "Experiment / N" ───────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(20)
    title_run = title_para.add_run(f"Experiment - {experiment_number}")
    _apply_tnr(title_run, 16, bold=True)

    # ── Aim ───────────────────────────────────────────────────────────────────
    _add_inline_section(doc, "Aim", aim)

    # ── Software Used ─────────────────────────────────────────────────────────
    _add_inline_section(doc, "Software Used", software)

    # ── Theory ────────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Theory")
    # Split on double newlines to get separate paragraphs
    theory_paras = [p.strip() for p in theory.strip().split("\n\n") if p.strip()]
    if theory_paras:
        for p in theory_paras:
            _add_body_paragraph(doc, p)
    else:
        _add_body_paragraph(doc, theory.strip())

    # ── Code ──────────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Code")
    _add_code_block(doc, code)

    # ── Output ────────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Output")
    
    # 1. Add terminal screenshot
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    img_para.paragraph_format.space_before = Pt(4)
    img_para.paragraph_format.space_after = Pt(8)
    img_run = img_para.add_run()
    img_run.add_picture(io.BytesIO(terminal_png), width=Inches(5.8))

    # 2. Add plot image if generated
    if plot_png:
        plot_para = doc.add_paragraph()
        plot_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        plot_para.paragraph_format.space_before = Pt(8)
        plot_para.paragraph_format.space_after = Pt(4)
        plot_run = plot_para.add_run()
        plot_run.add_picture(io.BytesIO(plot_png), width=Inches(5.8))

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
